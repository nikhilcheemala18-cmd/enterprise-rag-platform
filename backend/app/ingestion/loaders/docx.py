import hashlib
import re
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.ingestion.base import DocumentLoader
from app.models.document import (
    DocumentMetadata,
    DocxLocation,
    NormalizedDocument,
    SourceReference,
    TableElement,
    TextElement,
)

_HEADING_STYLE_RE = re.compile(r"^Heading (\d+)$")


def _heading_level(paragraph: Paragraph) -> int | None:
    style = paragraph.style
    if style is None or not style.name:
        return None
    match = _HEADING_STYLE_RE.match(style.name)
    if not match:
        return None
    return int(match.group(1))


class DocxLoader(DocumentLoader):
    """Loads a DOCX file into a NormalizedDocument.

    Paragraphs and tables are emitted as TextElement/TableElement in their
    original document order, by walking the body's XML children directly
    (python-docx's `.paragraphs` and `.tables` collections each lose the
    other's interleaving, so neither is used alone).

    Location convention: DocxLocation.paragraph_index and .table_index are
    zero-based counts of the paragraph/table's position among ALL
    paragraphs/tables in the document body, including paragraphs that were
    empty and therefore skipped -- so the index always reflects the true
    physical source location, not the position among emitted elements.
    section_index is left None: DOCX section breaks are not reliably
    mapped to the heading hierarchy in this implementation, and inventing
    that mapping is out of scope.

    Heading hierarchy: a simple stack of (level, text) tracks the current
    breadcrumb. On a heading paragraph, entries at the same or deeper
    level are popped before the new heading is pushed, and the heading's
    own section_path is the resulting breadcrumb (including itself). A
    non-heading paragraph or table gets the current breadcrumb as-is
    (not including any heading of its own, since it has none).
    """

    def load(
        self,
        file_path: str | Path,
        uploaded_by: str | None = None,
    ) -> NormalizedDocument:
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"DOCX file not found: {path}")
        if not path.is_file():
            raise ValueError(f"DOCX path is not a file: {path}")

        file_bytes = path.read_bytes()
        file_size = len(file_bytes)
        content_hash = hashlib.sha256(file_bytes).hexdigest()
        document_id = str(uuid.uuid4())

        try:
            document = DocxDocument(path)
        except Exception as exc:
            raise ValueError(f"Failed to open DOCX file {path}: {exc}") from exc

        metadata = DocumentMetadata(
            filename=path.name,
            file_type="docx",
            file_size=file_size,
            content_hash=content_hash,
            uploaded_by=uploaded_by,
        )

        elements: list[TextElement | TableElement] = []
        section_stack: list[tuple[int, str]] = []
        paragraph_position = 0
        table_position = 0

        body = document.element.body
        for child in body.iterchildren():
            if child.tag == qn("w:p"):
                paragraph = Paragraph(child, document)
                current_paragraph_index = paragraph_position
                paragraph_position += 1

                text = paragraph.text.strip()
                if not text:
                    continue

                level = _heading_level(paragraph)
                if level is not None:
                    while section_stack and section_stack[-1][0] >= level:
                        section_stack.pop()
                    section_stack.append((level, text))

                section_path = (
                    [t for _, t in section_stack] if section_stack else None
                )

                source = SourceReference(
                    document_id=document_id,
                    source_type="docx",
                    location=DocxLocation(
                        paragraph_index=current_paragraph_index,
                        table_index=None,
                        section_index=None,
                    ),
                )

                elements.append(
                    TextElement(
                        element_id=str(uuid.uuid4()),
                        element_index=len(elements),
                        source=source,
                        section_path=section_path,
                        content=text,
                        heading_level=level,
                    )
                )

            elif child.tag == qn("w:tbl"):
                table = Table(child, document)
                current_table_index = table_position
                table_position += 1

                rows_data = [[cell.text for cell in row.cells] for row in table.rows]
                columns = rows_data[0] if rows_data else []
                data_rows = rows_data[1:] if len(rows_data) > 1 else []

                section_path = (
                    [t for _, t in section_stack] if section_stack else None
                )

                source = SourceReference(
                    document_id=document_id,
                    source_type="docx",
                    location=DocxLocation(
                        paragraph_index=None,
                        table_index=current_table_index,
                        section_index=None,
                    ),
                )

                elements.append(
                    TableElement(
                        element_id=str(uuid.uuid4()),
                        element_index=len(elements),
                        source=source,
                        section_path=section_path,
                        columns=columns,
                        rows=data_rows,
                    )
                )

        return NormalizedDocument(
            document_id=document_id,
            metadata=metadata,
            elements=elements,
        )
