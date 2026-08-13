import shutil
import tempfile
import unittest
from pathlib import Path

from docx import Document as DocxDocument

from app.ingestion.base import DocumentLoader
from app.ingestion.loaders.docx import DocxLoader
from app.models.document import NormalizedDocument, TableElement, TextElement


def _build_fixture_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("Retention Strategy", level=1)
    doc.add_paragraph("Our retention program focuses on repeat purchases.")
    doc.add_paragraph("   ")  # whitespace-only, must be ignored
    doc.add_heading("Email Campaigns", level=2)
    doc.add_paragraph("Weekly campaigns are reviewed every Monday.")

    table = doc.add_table(rows=3, cols=2)
    table.cell(0, 0).text = "Metric"
    table.cell(0, 1).text = "Value"
    table.cell(1, 0).text = "Revenue"
    table.cell(1, 1).text = "$45,000"
    table.cell(2, 0).text = "Retention Rate"
    table.cell(2, 1).text = "42%"

    doc.add_paragraph("Campaign performance is reviewed monthly.")
    doc.save(path)


def _build_hierarchy_docx(path: Path) -> None:
    doc = DocxDocument()
    doc.add_heading("A", level=1)
    doc.add_heading("B", level=2)
    doc.add_heading("C", level=2)
    doc.add_heading("D", level=1)
    doc.save(path)


class TestDocxLoaderValidFixture(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp_dir = Path(tempfile.mkdtemp())
        cls.fixture_path = cls.tmp_dir / "fixture.docx"
        _build_fixture_docx(cls.fixture_path)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp_dir, ignore_errors=True)

    def setUp(self):
        self.loader = DocxLoader()
        self.document = self.loader.load(self.fixture_path, uploaded_by="user-1")

    def test_produces_normalized_document(self):
        self.assertIsInstance(self.document, NormalizedDocument)

    def test_document_id_generated(self):
        self.assertTrue(self.document.document_id)

    def test_metadata_filename(self):
        self.assertEqual(self.document.metadata.filename, "fixture.docx")

    def test_metadata_file_type(self):
        self.assertEqual(self.document.metadata.file_type, "docx")

    def test_metadata_file_size(self):
        self.assertGreater(self.document.metadata.file_size, 0)

    def test_uploaded_by_preserved(self):
        self.assertEqual(self.document.metadata.uploaded_by, "user-1")

    def test_content_hash_present(self):
        self.assertTrue(self.document.metadata.content_hash)
        self.assertEqual(len(self.document.metadata.content_hash), 64)

    def test_element_count_and_types(self):
        types = [type(e) for e in self.document.elements]
        self.assertEqual(
            types,
            [
                TextElement,
                TextElement,
                TextElement,
                TextElement,
                TableElement,
                TextElement,
            ],
        )

    def test_heading_becomes_text_element_with_level(self):
        heading1 = self.document.elements[0]
        self.assertEqual(heading1.content, "Retention Strategy")
        self.assertEqual(heading1.heading_level, 1)

        heading2 = self.document.elements[2]
        self.assertEqual(heading2.content, "Email Campaigns")
        self.assertEqual(heading2.heading_level, 2)

    def test_normal_paragraphs_become_text_element(self):
        para = self.document.elements[1]
        self.assertEqual(
            para.content, "Our retention program focuses on repeat purchases."
        )
        self.assertIsNone(para.heading_level)

    def test_empty_paragraphs_ignored(self):
        contents = [
            e.content for e in self.document.elements if isinstance(e, TextElement)
        ]
        self.assertNotIn("", contents)
        self.assertTrue(all(c.strip() for c in contents))

    def test_section_path_under_headings(self):
        para_under_h2 = self.document.elements[3]
        self.assertEqual(
            para_under_h2.section_path, ["Retention Strategy", "Email Campaigns"]
        )

    def test_heading_own_section_path_includes_itself(self):
        heading1 = self.document.elements[0]
        self.assertEqual(heading1.section_path, ["Retention Strategy"])

        heading2 = self.document.elements[2]
        self.assertEqual(
            heading2.section_path, ["Retention Strategy", "Email Campaigns"]
        )

    def test_table_becomes_table_element(self):
        table_element = self.document.elements[4]
        self.assertIsInstance(table_element, TableElement)

    def test_table_columns_preserved(self):
        table_element = self.document.elements[4]
        self.assertEqual(table_element.columns, ["Metric", "Value"])

    def test_table_rows_preserved(self):
        table_element = self.document.elements[4]
        self.assertEqual(
            table_element.rows,
            [["Revenue", "$45,000"], ["Retention Rate", "42%"]],
        )

    def test_table_section_path(self):
        table_element = self.document.elements[4]
        self.assertEqual(
            table_element.section_path, ["Retention Strategy", "Email Campaigns"]
        )

    def test_document_order_preserved(self):
        contents_or_kind = []
        for e in self.document.elements:
            if isinstance(e, TableElement):
                contents_or_kind.append("TABLE")
            else:
                contents_or_kind.append(e.content)
        self.assertEqual(
            contents_or_kind,
            [
                "Retention Strategy",
                "Our retention program focuses on repeat purchases.",
                "Email Campaigns",
                "Weekly campaigns are reviewed every Monday.",
                "TABLE",
                "Campaign performance is reviewed monthly.",
            ],
        )

    def test_element_index_sequential(self):
        indices = [e.element_index for e in self.document.elements]
        self.assertEqual(indices, [0, 1, 2, 3, 4, 5])

    def test_text_element_source_type(self):
        for e in self.document.elements:
            if isinstance(e, TextElement):
                self.assertEqual(e.source.source_type, "docx")

    def test_text_element_paragraph_index_populated(self):
        heading1 = self.document.elements[0]
        para1 = self.document.elements[1]
        heading2 = self.document.elements[2]
        para2 = self.document.elements[3]
        last_para = self.document.elements[5]

        self.assertEqual(heading1.source.location.paragraph_index, 0)
        self.assertEqual(para1.source.location.paragraph_index, 1)
        # index 2 is the skipped whitespace-only paragraph
        self.assertEqual(heading2.source.location.paragraph_index, 3)
        self.assertEqual(para2.source.location.paragraph_index, 4)
        self.assertEqual(last_para.source.location.paragraph_index, 5)
        self.assertIsNone(heading1.source.location.table_index)

    def test_table_element_table_index_populated(self):
        table_element = self.document.elements[4]
        self.assertEqual(table_element.source.location.table_index, 0)
        self.assertIsNone(table_element.source.location.paragraph_index)

    def test_source_document_id_matches(self):
        for e in self.document.elements:
            self.assertEqual(e.source.document_id, self.document.document_id)

    def test_serializes_to_json(self):
        payload = self.document.model_dump_json()
        self.assertIsInstance(payload, str)
        self.assertIn("Retention Strategy", payload)


class TestDocxLoaderHeadingHierarchy(unittest.TestCase):
    @classmethod
    def setUpClass(cls):
        cls.tmp_dir = Path(tempfile.mkdtemp())
        cls.fixture_path = cls.tmp_dir / "hierarchy.docx"
        _build_hierarchy_docx(cls.fixture_path)

    @classmethod
    def tearDownClass(cls):
        shutil.rmtree(cls.tmp_dir, ignore_errors=True)

    def test_heading_hierarchy_updates_correctly(self):
        loader = DocxLoader()
        document = loader.load(self.fixture_path)
        section_paths = [e.section_path for e in document.elements]
        self.assertEqual(
            section_paths,
            [
                ["A"],
                ["A", "B"],
                ["A", "C"],
                ["D"],
            ],
        )


class TestDocxLoaderEmptyDocument(unittest.TestCase):
    def test_empty_docx_returns_valid_document(self):
        tmp_dir = Path(tempfile.mkdtemp())
        try:
            path = tmp_dir / "empty.docx"
            DocxDocument().save(path)
            loader = DocxLoader()
            document = loader.load(path)
            self.assertIsInstance(document, NormalizedDocument)
            self.assertEqual(document.elements, [])
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)


class TestDocxLoaderMissingFile(unittest.TestCase):
    def test_missing_file_raises_clear_error(self):
        loader = DocxLoader()
        with self.assertRaises(FileNotFoundError):
            loader.load(Path(tempfile.gettempdir()) / "does_not_exist.docx")


class TestDocxLoaderInvalidFile(unittest.TestCase):
    def test_invalid_docx_raises_clear_error(self):
        tmp_dir = Path(tempfile.mkdtemp())
        try:
            path = tmp_dir / "invalid.docx"
            path.write_bytes(b"this is not a real docx file")
            loader = DocxLoader()
            with self.assertRaises(ValueError):
                loader.load(path)
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)


class TestDocxLoaderContract(unittest.TestCase):
    def test_loader_inherits_document_loader(self):
        self.assertIsInstance(DocxLoader(), DocumentLoader)


if __name__ == "__main__":
    unittest.main()
